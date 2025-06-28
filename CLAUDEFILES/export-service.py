# app/services/export_service.py
import io
import json
import tempfile
import os
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Pt, Inches
from ebooklib import epub
import html

class ExportService:
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def generate_pdf(self, story):
        """Generate a PDF from the story"""
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        premise_style = ParagraphStyle(
            'PremiseStyle',
            parent=styles['Italic'],
            fontSize=12,
            spaceBefore=0,
            spaceAfter=20,
            alignment=1  # Center alignment
        )
        
        chapter_title_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceBefore=20,
            spaceAfter=20,
            alignment=1  # Center alignment
        )
        
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontSize=12,
            spaceBefore=12,
            spaceAfter=12,
            leading=14
        )
        
        # Create content elements
        elements = []
        
        # Add title and premise
        elements.append(Paragraph(story.title, title_style))
        if story.premise:
            elements.append(Paragraph(story.premise, premise_style))
        elements.append(Spacer(1, 30))
        
        # Get chapters
        metadata = json.loads(story.metadata) if story.metadata else {}
        
        # Add chapters
        for chapter in story.chapters:
            elements.append(PageBreak())
            elements.append(Paragraph(chapter.title, chapter_title_style))
            
            # Process chapter content
            paragraphs = chapter.content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    elements.append(Paragraph(para, body_style))
        
        # Add metadata page at the end
        elements.append(PageBreak())
        elements.append(Paragraph("About This Story", title_style))
        
        elements.append(Paragraph(f"<b>Genre:</b> {metadata.get('genre', 'Not specified')}", body_style))
        elements.append(Paragraph(f"<b>Theme:</b> {metadata.get('theme', 'Not specified')}", body_style))
        elements.append(Paragraph(f"<b>Target Audience:</b> {metadata.get('targetAudience', 'Not specified')}", body_style))
        elements.append(Paragraph(f"<b>Tone:</b> {metadata.get('tone', 'Not specified')}", body_style))
        
        if metadata.get('uniqueElements'):
            elements.append(Paragraph("<b>Unique Elements:</b>", body_style))
            for element in metadata.get('uniqueElements', []):
                elements.append(Paragraph(f"• {element}", body_style))
        
        if metadata.get('keySymbols'):
            elements.append(Paragraph("<b>Key Symbols:</b>", body_style))
            for symbol in metadata.get('keySymbols', []):
                elements.append(Paragraph(f"• {symbol}", body_style))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer
    
    def generate_docx(self, story):
        """Generate a DOCX from the story"""
        # Create DOCX document
        doc = Document()
        
        # Add title and premise
        doc.add_heading(story.title, 0)
        if story.premise:
            premise_paragraph = doc.add_paragraph()
            premise_run = premise_paragraph.add_run(story.premise)
            premise_run.italic = True
            premise_paragraph.alignment = 1  # Center alignment
        
        doc.add_page_break()
        
        # Get metadata
        metadata = json.loads(story.metadata) if story.metadata else {}
        
        # Add chapters
        for chapter in story.chapters:
            doc.add_heading(chapter.title, 1)
            
            # Process chapter content
            paragraphs = chapter.content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)
            
            doc.add_page_break()
        
        # Add metadata page at the end
        doc.add_heading("About This Story", 1)
        
        doc.add_paragraph(f"Genre: {metadata.get('genre', 'Not specified')}")
        doc.add_paragraph(f"Theme: {metadata.get('theme', 'Not specified')}")
        doc.add_paragraph(f"Target Audience: {metadata.get('targetAudience', 'Not specified')}")
        doc.add_paragraph(f"Tone: {metadata.get('tone', 'Not specified')}")
        
        if metadata.get('uniqueElements'):
            doc.add_paragraph("Unique Elements:")
            for element in metadata.get('uniqueElements', []):
                doc.add_paragraph(f"• {element}", style='ListBullet')
        
        if metadata.get('keySymbols'):
            doc.add_paragraph("Key Symbols:")
            for symbol in metadata.get('keySymbols', []):
                doc.add_paragraph(f"• {symbol}", style='ListBullet')
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generate_epub(self, story):
        """Generate an EPUB from the story"""
        # Create EPUB book
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier(f'storyforge-{story.id}')
        book.set_title(story.title)
        book.set_language('en')
        
        metadata = json.loads(story.metadata) if story.metadata else {}
        if metadata.get('genre'):
            book.add_metadata('DC', 'subject', metadata.get('genre'))
        
        # Add CSS
        style = '''
        @namespace epub "http://www.idpf.org/2007/ops";
        body {
            font-family: Cambria, Liberation Serif, Bitstream Vera Serif, Georgia, Times, Times New Roman, serif;
            margin: 5%;
            text-align: justify;
        }
        h1, h2 {
            text-align: center;
            page-break-before: always;
        }
        .title {
            margin: 3em 0;
            text-align: center;
        }
        .premise {
            margin: 1em 10%;
            text-align: center;
            font-style: italic;
        }
        '''
        
        css = epub.EpubItem(
            uid="style_default",
            file_name="style/default.css",
            media_type="text/css",
            content=style
        )
        book.add_item(css)
        
        # Add title page
        title_page = epub.EpubHtml(title='Title Page', file_name='title_page.xhtml')
        title_page.content = f'''
        <html xmlns="http://www.w3.org/1999/xhtml">
        <head>
            <title>{story.title}</title>
            <link rel="stylesheet" href="style/default.css" type="text/css" />
        </head>
        <body>
            <div class="title">
                <h1>{story.title}</h1>
            </div>
            
            {'<div class="premise"><p>' + story.premise + '</p></div>' if story.premise else ''}
        </body>
        </html>
        '''
        book.add_item(title_page)
        
        # Add chapters
        chapters = []
        for i, chapter in enumerate(story.chapters):
            # Create chapter
            c = epub.EpubHtml(
                title=chapter.title,
                file_name=f'chapter_{i+1}.xhtml',
                lang='en'
            )
            
            # Process chapter content
            paragraphs = chapter.content.split('\n\n')
            formatted_paragraphs = [f'<p>{para}</p>' for para in paragraphs if para.strip()]
            
            c.content = f'''
            <html xmlns="http://www.w3.org/1999/xhtml">
            <head>
                <title>{chapter.title}</title>
                <link rel="stylesheet" href="style/default.css" type="text/css" />
            </head>
            <body>
                <h2>{chapter.title}</h2>
                {''.join(formatted_paragraphs)}
            </body>
            </html>
            '''
            
            book.add_item(c)
            chapters.append(c)
        
        # Add metadata page
        about_page = epub.EpubHtml(title='About This Story', file_name='about.xhtml')
        about_page.content = f'''
        <html xmlns="http://www.w3.org/1999/xhtml">
        <head>
            <title>About This Story</title>
            <link rel="stylesheet" href="style/default.css" type="text/css" />
        </head>
        <body>
            <h2>About This Story</h2>
            <p><strong>Genre:</strong> {metadata.get('genre', 'Not specified')}</p>
            <p><strong>Theme:</strong> {metadata.get('theme', 'Not specified')}</p>
            <p><strong>Target Audience:</strong> {metadata.get('targetAudience', 'Not specified')}</p>
            <p><strong>Tone:</strong> {metadata.get('tone', 'Not specified')}</p>
            
            {('<p><strong>Unique Elements:</strong></p><ul>' + 
              ''.join([f'<li>{element}</li>' for element in metadata.get('uniqueElements', [])]) + 
              '</ul>') if metadata.get('uniqueElements') else ''}
              
            {('<p><strong>Key Symbols:</strong></p><ul>' + 
              ''.join([f'<li>{symbol}</li>' for symbol in metadata.get('keySymbols', [])]) + 
              '</ul>') if metadata.get('keySymbols') else ''}
        </body>
        </html>
        '''
        book.add_item(about_page)
        chapters.append(about_page)
        
        # Define Table Of Contents
        book.toc = (
            (epub.Section('Title Page'), (title_page,)),
            (epub.Section('Chapters'), chapters)
        )
        
        # Add default NCX and Nav
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Define spine
        book.spine = ['nav', title_page] + chapters
        
        # Create the EPUB file
        buffer = io.BytesIO()
        epub.write_epub(buffer, book)
        buffer.seek(0)
        return buffer

# Create singleton instance
export_service = ExportService()

# For the generate_pdf, generate_docx, and generate_epub functions in the story_bp routes,
# replace the function calls with:
# from app.services.export_service import export_service
# ...
# file_stream = export_service.generate_pdf(story)
# file_stream = export_service.generate_docx(story)
# file_stream = export_service.generate_epub(story)
